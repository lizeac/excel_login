from openpyxl import load_workbook, Workbook
from datetime import datetime
from collections import Counter
from docx import Document
from win32com import client


dates = datetime.now()
year = dates.strftime('%Y')
month = dates.strftime('%m')
user_access_date = dates.strftime('%d/%m/%Y')
user_access_hour = dates.strftime('%H:%M')
file_to_handle = f'users_{year}.xlsx'



wb = load_workbook(file_to_handle)
ws = wb.active  
month_name = {
            '1': 'Janeiro',
            '2': 'Fevereiro',
            '3': 'Março',
            '4': 'Abril',
            '5': 'Maio',
            '6': 'Junho',
            '7': 'Julho',
            '8': 'Agosto',
            '9': 'Setembro',
            '10': 'Outubro',
            '11': 'Novembro',
            '12': 'Dezembro'        
        }

file_month = f'{month_name[str(int(month))]}'

# primeira tentativa da função

# função pra extrair dados das colunas no excel
def cell_values(letra: str):
    return [cell.value for cell in ws[letra][1:]]


general_data = {
    'Matricula': [*cell_values('A')],
    'Nome': [*cell_values('B')],
    'Servico': [*cell_values('C')], 
    'Data de Acesso': [*cell_values('D')], 
    'Hora de Acesso': [*cell_values('E')],
    'Visitantes?' : [*cell_values('F')]
}
# Função pra gerar o item mais logico de cada mes.
# Achei necessário fazer dessa forma pois existe a possibilidade de empate na primeira colocação
# se utilizar só o counter 'cru' ele pega um valor só
def most_often(lista):
    counter = Counter(lista).most_common()
    most_frequent = [item[0] for item in counter if item[1] == counter[0][1]]
    quantity = counter[0][1]
    return most_frequent, quantity

total_visitantes = int()
for item in general_data['Visitantes?']:
    if 'Usuário UFBA' not in item:
        total_visitantes += 1
        print(item)
print(total_visitantes)
# primeiro teste da função pra pegar os mais frequentes
# def most_often( lista):
#     result = []
#     counter = Counter(lista).most_common()
#     key = 0
#     quantity = counter[0][1]
#     for item in counter:
#         if quantity in item:
#             result.append(item[0])
#     print(result, quantity)
#     return result, quantity


name_most = most_often(general_data['Nome'])
# names = 
mat_most = most_often(general_data['Matricula'])
day_most = most_often(general_data['Data de Acesso'])
# documento = Document(f'.\{file_month.capitalize()}.docx')

documento = Document()
serv = most_often(general_data['Servico'])
visit = most_often(general_data['Visitantes?'])


documento.add_heading(f'Relatório de Usuários da Biblioteca do Mês de {month_name[str(int(month))].capitalize()}', 0)
documento.add_heading('O Nome dos(as) visitantes mais frequente é: ', 3)
paragrafo = documento.add_paragraph().add_run(
    f'''{name_most[0]}. 
A quantidade de visitas de cada à biblioteca foi de {name_most[1]}'''
)
documento.add_heading(
    'Os números de matrícula(s) correspondente(s), respectivamente: ', 3
)
documento.add_paragraph().add_run(
    f'{mat_most[0]}'
)
documento.add_heading(
    'O Serviço mais utilizado foi de: ', 3
)
documento.add_paragraph().add_run(
    f'''{serv[0]}. 
Com um total de {serv[1]} utilizações.'''
)
documento.add_heading(
    'O dia que mais recebeu visitas foi:', 3
)
documento.add_paragraph().add_run(
    f'''{day_most[0]}. 
Com um total de {day_most[1]} visitas.'''
)
documento.add_heading(
    'A quantidade total de visitantes do mês foi de ', 3
)
documento.add_paragraph().add_run(
    f'''{len(general_data["Nome"])} pessoas.'''
)
documento.add_heading(
    'Neste mês, o total de visitantes foi de: ', 3
)
documento.add_paragraph().add_run(
    f'''{total_visitantes}'''
)

documento.add_paragraph('Fim do Relatório')

# ATIVAR DEPOIS PRA FAZER OS TESTES
# documento.save(f'Relatorio_{file_month.capitalize()}.docx')

documento.save(f'Relatorio_{month_name[str(int(month))]}.docx')


# Transformar o arquivo word em Excel
wdFormatPDF = 17

entrada = f'excel_login/Relatorio_{month_name[str(int(month))]}.docx'
saida = f'excel_login/Relatorio_{month_name[str(int(month))]}.pdf'

word = client.Dispatch('Word.Application')
doc = word.Documents.Open(entrada)
doc.SaveAs(saida, FileFormat=wdFormatPDF)
doc.Close()
word.Quit()